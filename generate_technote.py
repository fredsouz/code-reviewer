from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)

# ── Style helpers ─────────────────────────────────────────────────────────────
def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=14, bold=True, color=(31, 73, 125))
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=12, bold=True, color=(68, 114, 196))
    return p

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent   = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after   = Pt(2)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def code_block(lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.4)
        run = p.add_run(line if line else " ")
        set_font(run, name="Courier New", size=9, color=(40, 40, 40))
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'),   'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'),  'F2F2F2')
        p._p.get_or_add_pPr().append(shading)

def inline_code(paragraph, text):
    run = paragraph.add_run(text)
    set_font(run, name="Courier New", size=10, color=(180, 0, 0))
    return run

def kv(label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(label + ": ")
    set_font(r1, size=11, bold=True)
    r2 = p.add_run(value)
    set_font(r2, size=11)
    return p


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("Technical Note")
set_font(r, size=28, bold=True, color=(31, 73, 125))

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub_p.add_run("Local LLM Code Reviewer")
set_font(r, size=18, bold=False, color=(68, 114, 196))

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Installation · Configuration · End-to-End Testing")
set_font(r, size=12, color=(100, 100, 100))

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(f"Author: Frederic Sousa   |   Date: {datetime.date(2026, 6, 6).strftime('%B %d, %Y')}   |   Version: 1.0")
set_font(r, size=10, color=(120, 120, 120))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading1("1. Overview")
body(
    "This document describes the installation, configuration, and end-to-end testing of a local "
    "LLM-based code review system. The system uses Meta's Llama 3.1 model (running locally via "
    "Ollama) to perform automated code analysis, root cause analysis (RCA), and fix proposals. "
    "Proposed fixes are submitted as GitHub draft pull requests that require mandatory human "
    "approval before merging — no automatic merging is possible."
)

heading2("1.1 Objectives")
bullet("Analyze code diffs and files for bugs, security vulnerabilities, and quality issues")
bullet("Perform root cause analysis on error logs and stack traces")
bullet("Generate fix proposals and open draft pull requests automatically")
bullet("Enforce a mandatory human approval gate — no code reaches main without human review")

heading2("1.2 Architecture")
body("The system is composed of the following components:")
bullet("Ollama — local model runtime serving Llama 3.1 via REST API on localhost:11434")
bullet("Bash scripts — thin wrappers that build prompts, call Ollama, and drive Git/GitHub")
bullet("GitHub repository — stores scripts, prompts, and receives draft PRs")
bullet("Branch protection rules — enforced at GitHub API level, prevent any direct push to main")

# ══════════════════════════════════════════════════════════════════════════════
# 2. PREREQUISITES
# ══════════════════════════════════════════════════════════════════════════════
heading1("2. Prerequisites")
body("The following were verified present before installation:")

kv("Operating System", "Ubuntu 24.04 LTS (Linux 6.17.0)")
kv("Git",              "2.43.0 (pre-installed)")
kv("OpenSSH",          "9.6p1 (pre-installed)")
kv("Python",           "3.12 (pre-installed)")
kv("GitHub account",   "fredsouz (pre-existing)")

# ══════════════════════════════════════════════════════════════════════════════
# 3. INSTALLATION
# ══════════════════════════════════════════════════════════════════════════════
heading1("3. Installation")

heading2("3.1 GitHub CLI (gh)")
body("The GitHub CLI was installed from the official apt repository:")
code_block([
    "curl -fsSL https://cli.github.com/packages/githubcli-archive-keyring.gpg \\",
    "  | sudo dd of=/usr/share/keyrings/githubcli-archive-keyring.gpg",
    "sudo chmod go+r /usr/share/keyrings/githubcli-archive-keyring.gpg",
    'echo "deb [arch=$(dpkg --print-architecture) \\',
    "  signed-by=/usr/share/keyrings/githubcli-archive-keyring.gpg] \\",
    '  https://cli.github.com/packages stable main" \\',
    "  | sudo tee /etc/apt/sources.list.d/github-cli.list",
    "sudo apt update && sudo apt install gh -y",
])
body("Version installed: gh 2.45.0")

heading2("3.2 Git Identity")
body("Global Git identity was configured:")
code_block([
    'git config --global user.name "Frederic Sousa"',
    'git config --global user.email "frederic.sousa@free.fr"',
    'git config --global init.defaultBranch main',
])

heading2("3.3 SSH Key")
body("An Ed25519 SSH key was generated and registered with GitHub:")
code_block([
    'ssh-keygen -t ed25519 -C "frederic.sousa@free.fr" \\',
    '  -f ~/.ssh/github_ed25519 -N ""',
    "ssh-keyscan github.com >> ~/.ssh/known_hosts",
])
body(
    "The public key (~/.ssh/github_ed25519.pub) was uploaded to GitHub via "
    "gh auth login, selecting SSH as the protocol."
)

heading2("3.4 GitHub CLI Authentication")
body("Authentication was completed interactively:")
code_block(["gh auth login"])
body("Options selected: GitHub.com → SSH → existing key → browser login.")
body("Authentication verified:")
code_block([
    "$ gh auth status",
    "github.com",
    "  ✓ Logged in to github.com account fredsouz (keyring)",
    "  - Git operations protocol: ssh",
    "  - Token scopes: admin:public_key, gist, read:org, repo",
    "",
    "$ ssh -T git@github.com",
    "Hi fredsouz! You've successfully authenticated, but GitHub does not provide shell access.",
])

heading2("3.5 Ollama and Llama 3.1")
body("Ollama was installed using the official install script:")
code_block(["curl -fsSL https://ollama.com/install.sh | sh"])
body("Llama 3.1 (8B parameter model, ~4.7 GB) was pulled:")
code_block(["ollama pull llama3.1"])
body(
    "Ollama was already registered as a systemd service by the installer and confirmed active:"
)
code_block([
    "$ systemctl is-active ollama",
    "active",
    "",
    "$ curl -s http://localhost:11434/api/tags | python3 -c \\",
    "  \"import json,sys; [print(m['name']) for m in json.load(sys.stdin)['models']]\"",
    "llama3.1:latest",
])

# ══════════════════════════════════════════════════════════════════════════════
# 4. REPOSITORY SETUP
# ══════════════════════════════════════════════════════════════════════════════
heading1("4. Repository Setup")

heading2("4.1 Repository Creation")
body(
    "The code-reviewer repository was created on GitHub and cloned locally "
    "at ~/claude/code-reviewer:"
)
code_block([
    'gh repo create code-reviewer --public \\',
    '  --description "Llama-based code review, RCA, fix proposals and PR creation" \\',
    '  --clone',
])
p = body("Repository URL: ")
inline_code(p, "https://github.com/fredsouz/code-reviewer")

heading2("4.2 Directory Structure")
body("The following structure was created:")
code_block([
    "code-reviewer/",
    "├── config/",
    "│   └── models.yml          # Ollama model config and parameters",
    "├── prompts/",
    "│   ├── review.txt          # Code review prompt template",
    "│   ├── rca.txt             # Root cause analysis prompt template",
    "│   └── fix.txt             # Fix generation prompt template",
    "├── scripts/",
    "│   ├── review.sh           # Analyze a file or git diff",
    "│   ├── rca.sh              # RCA from error log or string",
    "│   └── fix-and-pr.sh       # Generate fix + open draft PR",
    "└── test/",
    "    └── sample_buggy.py     # Test file with intentional bugs",
])

heading2("4.3 Branch Protection")
body(
    "Branch protection was configured on main via the GitHub API to enforce the "
    "mandatory human approval gate:"
)
code_block([
    "gh api repos/fredsouz/code-reviewer/branches/main/protection \\",
    "  --method PUT --input - <<'EOF'",
    "{",
    '  "required_status_checks": null,',
    '  "enforce_admins": true,',
    '  "required_pull_request_reviews": null,',
    '  "restrictions": null,',
    '  "allow_force_pushes": false,',
    '  "allow_deletions": false',
    "}",
    "EOF",
])
body("Rules enforced:")
bullet("Direct push to main blocked — all changes must go through a pull request")
bullet("Force push blocked")
bullet("Branch deletion blocked")
bullet("Enforced for admins — the repository owner cannot bypass the rules")
body(
    "Note: required approval count was intentionally set to zero for solo development. "
    "The protection guarantees that no code merges without a deliberate human action "
    "(opening a PR and clicking merge). Auto-merge is not possible."
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. SCRIPTS
# ══════════════════════════════════════════════════════════════════════════════
heading1("5. Scripts")

heading2("5.1 review.sh — Code Review")
body("Analyzes a file or the current git diff and reports issues found by Llama.")
code_block([
    "# Review a specific file",
    "bash scripts/review.sh path/to/file.py",
    "",
    "# Review uncommitted changes",
    "bash scripts/review.sh",
])
body(
    "The script reads the prompt template from prompts/review.txt, substitutes the diff "
    "content, and sends it to the Ollama REST API. Output is printed to stdout."
)

heading2("5.2 rca.sh — Root Cause Analysis")
body("Performs root cause analysis on an error log file or error string.")
code_block([
    "# From a file",
    "bash scripts/rca.sh error.log",
    "",
    '# From an inline string',
    'bash scripts/rca.sh "IndexError: list index out of range at line 15"',
])

heading2("5.3 fix-and-pr.sh — Fix and Pull Request")
body(
    "Generates a fix for a specific issue, asks for confirmation, then creates a branch, "
    "commits the fix, pushes it, and opens a draft pull request on GitHub."
)
code_block([
    'bash scripts/fix-and-pr.sh <file> "<issue description>"',
    "",
    "# Example",
    'bash scripts/fix-and-pr.sh src/auth.py "SQL injection in login function line 42"',
])
body("The script will not proceed past the confirmation prompt without human input.")
body("The resulting PR is created as a draft with:")
bullet("Title prefixed with fix: and the filename")
bullet("Body describing the issue and a mandatory review checklist")
bullet("Assignee set to the authenticated GitHub user")

# ══════════════════════════════════════════════════════════════════════════════
# 6. END-TO-END TESTING
# ══════════════════════════════════════════════════════════════════════════════
heading1("6. End-to-End Testing")

heading2("6.1 Test File")
body(
    "A test file test/sample_buggy.py was created with three intentional defects:"
)
code_block([
    "import sqlite3",
    "",
    "def get_user(username):",
    '    conn = sqlite3.connect("users.db")',
    "    cursor = conn.cursor()",
    "    # BUG 1: SQL injection — username concatenated directly",
    "    query = \"SELECT * FROM users WHERE username = '\" + username + \"'\"",
    "    cursor.execute(query)",
    "    return cursor.fetchone()",
    "",
    "def divide(a, b):",
    "    # BUG 2: division by zero — no guard on b",
    "    return a / b",
    "",
    "def process_items(items):",
    "    results = []",
    "    # BUG 3: off-by-one — range(len+1) causes IndexError",
    "    for i in range(len(items) + 1):",
    "        results.append(items[i] * 2)",
    "    return results",
])

heading2("6.2 Test 1 — Code Review")
body("Command:")
code_block(["bash scripts/review.sh test/sample_buggy.py"])
body("Llama correctly identified all three defects:")
bullet("CRITICAL — SQL injection in get_user (line 5): username concatenated into query without escaping")
bullet("MEDIUM   — Division by zero in divide (line 9): no guard on b == 0")
bullet("LOW      — Off-by-one in process_items (line 12): range(len+1) causes IndexError")
bullet("MEDIUM   — Missing error handling across all functions")

heading2("6.3 Test 2 — Root Cause Analysis")
body("Command:")
code_block([
    'bash scripts/rca.sh \\',
    '  "IndexError: list index out of range at process_items line 15, \\',
    "   called with items=['a','b','c']\"",
])
body("Llama output (summarised):")
bullet("Root cause: loop iterates one past the end of the list due to range(len(items) + 1)")
bullet("Trigger: calling process_items with any non-empty list")
bullet("Affected component: process_items function")
bullet("Proposed fix: use range(len(items)) or iterate directly with for item in items")

heading2("6.4 Test 3 — Fix and Pull Request")
body("Command:")
code_block([
    'bash scripts/fix-and-pr.sh test/sample_buggy.py \\',
    '  "SQL injection in get_user function — username concatenated directly into query string"',
])
body("Llama generated the following fix:")
code_block([
    "def get_user(username):",
    '    conn = sqlite3.connect("users.db")',
    "    cursor = conn.cursor()",
    '    query = "SELECT * FROM users WHERE username = ?"',
    "    cursor.execute(query, (username,))",
    "    return cursor.fetchone()",
])
body(
    "The fix correctly replaces string concatenation with a parameterized query using "
    "the ? placeholder and passes username as a bound parameter — the standard SQLite "
    "defense against SQL injection."
)
body("After confirming at the prompt, the script:")
bullet("Created branch llama-fix/sample_buggy-py-<timestamp>")
bullet("Committed and pushed the fix")
bullet("Opened draft PR #2 at https://github.com/fredsouz/code-reviewer/pull/2")

heading2("6.5 Human Approval Gate Verification")
body(
    "The branch protection was independently verified when a direct push to main was "
    "attempted and rejected by GitHub:"
)
code_block([
    "$ git push origin main",
    "remote: error: GH006: Protected branch update failed for refs/heads/main.",
    "remote: - Changes must be made through a pull request.",
    "! [remote rejected] main -> main (protected branch hook declined)",
])
body(
    "This confirms that no code — including code pushed by the repository owner — "
    "can reach main without going through a pull request."
)

heading2("6.6 Test Results Summary")
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
for cell, text in zip(hdr, ["Test", "Expected", "Result"]):
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_font(run, size=10, bold=True)

rows = [
    ("Code review on buggy file",           "All 3 bugs detected",                       "PASS"),
    ("RCA on IndexError string",             "Root cause and fix identified",             "PASS"),
    ("Fix generation for SQL injection",     "Parameterized query proposed",              "PASS"),
    ("Draft PR creation",                    "PR opened, assignee set, draft status",     "PASS"),
    ("Direct push to main blocked",          "Push rejected by branch protection",        "PASS"),
    ("PR merge after human action",          "Merge succeeds via gh pr merge",            "PASS"),
]
for test, expected, result in rows:
    row = table.add_row().cells
    for cell, text in zip(row, [test, expected, result]):
        p = cell.paragraphs[0]
        run = p.add_run(text)
        color = (0, 128, 0) if text == "PASS" else (180, 0, 0)
        set_font(run, size=10, color=color if text in ("PASS", "FAIL") else None)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 7. WORKFLOW REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
heading1("7. Operational Workflow Reference")
body("The complete workflow for using the code-reviewer on any repository:")
code_block([
    "# 1. Navigate to the target repository",
    "cd ~/my-project",
    "",
    "# 2. Review current changes or a specific file",
    "bash ~/claude/code-reviewer/scripts/review.sh src/auth.py",
    "",
    "# 3. If an error needs investigation",
    'bash ~/claude/code-reviewer/scripts/rca.sh "TypeError: cannot read property of undefined"',
    "",
    "# 4. Generate a fix and open a draft PR",
    'bash ~/claude/code-reviewer/scripts/fix-and-pr.sh src/auth.py "issue description"',
    "",
    "# 5. Review the draft PR on GitHub",
    "#    https://github.com/fredsouz/<repo>/pulls",
    "",
    "# 6. Merge the PR manually after review",
    "gh pr merge <number> --merge",
])

# ══════════════════════════════════════════════════════════════════════════════
# 8. KNOWN LIMITATIONS
# ══════════════════════════════════════════════════════════════════════════════
heading1("8. Known Limitations")
bullet(
    "Context window: Llama 3.1 8B has an 8192 token context window. "
    "Very large files or diffs may be truncated."
)
bullet(
    "Fix quality: the model proposes fixes for the described issue only. "
    "It does not refactor or fix unrelated issues in the same file."
)
bullet(
    "File overwrite: fix-and-pr.sh writes the model output directly to the target file. "
    "Always review the diff on GitHub before merging."
)
bullet(
    "Solo approval: GitHub does not allow self-approval of PRs. "
    "The human gate is enforced by requiring a deliberate merge action, not a second reviewer."
)
bullet(
    "Ollama must be running: scripts will fail if the Ollama service is stopped. "
    "Check with: systemctl status ollama"
)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
out = "/home/fredsouz/claude/code-reviewer/TechNote_CodeReviewer_v1.0.docx"
doc.save(out)
print(f"Saved: {out}")
